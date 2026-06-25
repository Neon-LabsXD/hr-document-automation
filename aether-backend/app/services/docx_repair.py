from io import BytesIO

from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx import Document


def _find_template_spans(text: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    cursor = 0

    while cursor < len(text):
        start = text.find("{{", cursor)
        if start == -1:
            break

        end = text.find("}}", start + 2)
        if end == -1:
            break

        spans.append((start, end + 2))
        cursor = end + 2

    return spans


def _repair_paragraph(paragraph: Paragraph) -> None:
    runs = list(paragraph.runs)

    if len(runs) < 2:
        return

    full_text = "".join(run.text for run in runs)

    if "{{" not in full_text or "}}" not in full_text:
        return

    char_run_indexes: list[int] = []
    for run_index, run in enumerate(runs):
        char_run_indexes.extend([run_index] * len(run.text))

    if not char_run_indexes:
        return

    target_by_char: dict[int, int] = {}
    for start, end in _find_template_spans(full_text):
        affected_runs = set(char_run_indexes[start:end])

        if len(affected_runs) <= 1:
            continue

        target_run_index = min(affected_runs)
        for char_index in range(start, end):
            target_by_char[char_index] = target_run_index

    if not target_by_char:
        return

    repaired_text_by_run = [""] * len(runs)
    for char_index, char in enumerate(full_text):
        run_index = target_by_char.get(char_index, char_run_indexes[char_index])
        repaired_text_by_run[run_index] += char

    for run, repaired_text in zip(runs, repaired_text_by_run):
        run.text = repaired_text


def _iter_block_items(parent: DocumentObject | Table):
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
    else:
        parent_element = parent._tbl

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _repair_table(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _repair_paragraph(paragraph)
            for nested_table in cell.tables:
                _repair_table(nested_table)


def repair_docx_template(template_content: bytes) -> bytes:
    """
    Joins Word runs that split Jinja variables like {{candidate_name}}.
    This keeps text outside template tags in its original runs where possible.
    """
    document = Document(BytesIO(template_content))

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            _repair_paragraph(block)
        elif isinstance(block, Table):
            _repair_table(block)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
